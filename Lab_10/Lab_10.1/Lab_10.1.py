from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH # нужен для того, чтобы выравнивать текст в таблице по центру
from docx.oxml import OxmlElement     #|
from docx.oxml.ns import qn           #| нужно для окрашивания фона ячейки таблицы

document = Document() # создаем чистый документ Word

document.add_paragraph("В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти:") # добавляем первый параграф

document.add_paragraph("Флеш-память: используется для хранения скетчей.", style="List Bullet 3")  # добавляем первый элемент списка с двойным отступом
record2 = document.add_paragraph("ОЗУ (", style="List Bullet 3") # добавляем начало второго элемента списка
record2.add_run("SRAM ").bold = True # добавляем текст с жирным шрифтом
record2.add_run("- static random access memory").italic = True # добавляем текст с курсивом
record2.add_run(", статическая оперативная память с произвольным доступом): используется для хранения постоянной информации.") # добавляем оставшуюся часть второго элемента списка
document.add_paragraph("EEPROM (энергозависимая память): используется для хранения постоянной информации.", style="List Bullet 3") # добавляем третий элемент списка

document.add_paragraph("Флеш-память и EEPROM являются энергонезависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью.") # добавляем второй параграф

table = document.add_table(rows=4, cols=5) # создаем пустую таблицу размером 4х5
table.style = "Table Grid"
contents = (
        ("", "ATmega168", "ATmega328", "ATmega1280", "ATmega2560"),
        ("Flash (1 кБ flash-памяти занят загрузчиком)", "16 Кбайт", "32 Кбайт", "128 Кбайт", "256 Кбайт"),
        ("SRAM", "1 Кбайт", "2 Кбайт", "8 Кбайт", "8 Кбайт"),
        ("EEPROM", "512 байт", "1024 байта", "4 Кбайт", "4 Кбайт")
        ) # содержание таблицы
headers = table.rows[0].cells # выбираем ячейки первой строки таблицы, в которой будут заголовки столбцов
for i in range(5): # добавляем заголовки столбцов таблицы с жирным шрифтом и выравниванием по центру
    cell = headers[i]
    cellXmlElement = cell._tc                                   #|
    cellProperties = cellXmlElement.get_or_add_tcPr()           #|
    shadeObj = OxmlElement("w:shd")                             #| окрашиваем фон ячейки в серый цвет
    shadeObj.set(qn("w:fill"), "bdd6ee")                        #|
    cellProperties.append(shadeObj)                             #|
    header = headers[i].add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run(contents[0][i]).bold = True
entries = table.columns[0].cells # выбираем ячейки первого столбца таблицы, в которой будут заголовки строк
for i in range(1, 4): # добавляем заголовки строк таблицы с жирным шрифтом и выравниванием по центру
    cell = entries[i]
    cellXmlElement = cell._tc                                   #|
    cellProperties = cellXmlElement.get_or_add_tcPr()           #|
    shadeObj = OxmlElement("w:shd")                             #| окрашиваем фон ячейки в серый цвет
    shadeObj.set(qn("w:fill"), "bdd6ee")                        #|
    cellProperties.append(shadeObj)                             #|
    entry = entries[i].add_paragraph()
    entry.alignment = WD_ALIGN_PARAGRAPH.CENTER
    entry.add_run(contents[i][0]).bold = True
for row in range(1, 4): # заполняем таблицу содержимым (выравнивание также по центру)
    for col in range(1, 5):
        data = table.cell(row, col).add_paragraph()
        data.alignment = WD_ALIGN_PARAGRAPH.CENTER
        data.add_run(contents[row][col])
    
document.add_paragraph() # не уверен пустая ли это строка в задании, или это отступ от таблицы, созданный автоматически

lastParagraph = document.add_paragraph() # добавляем последний параграф с курсивом
lastParagraph.add_run("Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25°C. Эти данные не распространяются на операции чтения данных из EEPROM — чтение данных не лимитированно. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM.").italic = True

document.save("Word_1.docx") # сохраняем документ