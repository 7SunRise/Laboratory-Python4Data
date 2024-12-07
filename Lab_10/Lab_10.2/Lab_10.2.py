from docx import Document
from docx.shared import Cm

document = Document("Word_1.docx") # загружаем текстовый файл, в который будет добавляться картинка

document.add_picture("Picture.png", width=Cm(5), height=Cm(5)) # добавляем картинку с заданной высотой и шириной
document.add_paragraph("Подпись: Новогоднее настроение") # добавляем подпись
    
document.save("Word_2.docx") # сохраняем файл